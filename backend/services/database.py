import hashlib
import json
import logging
import os
from collections.abc import Iterator

from docx import Document as DocxDocumentFile
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma

from backend.config import get_settings
from backend.services.llm import get_embedding_model

logger = logging.getLogger(__name__)

INDEX_SCHEMA_VERSION = "v3_section_chunking"
MIN_CHUNK_CHARS = 180


def _get_doc_path() -> str:
    settings = get_settings()
    return os.path.join(settings.data_dir, settings.doc_filename)


def _normalize_source_path(path: str) -> str:
    return os.path.abspath(path)


def _compute_source_hash(path: str) -> str:
    with open(path, "rb") as f:
        return hashlib.sha256(f.read()).hexdigest()


def get_vectorstore() -> Chroma:
    settings = get_settings()
    return Chroma(
        collection_name=settings.collection_name,
        embedding_function=get_embedding_model(),
        persist_directory=settings.chroma_persist_dir,
    )


def collection_is_empty(vectorstore: Chroma) -> bool:
    return vectorstore._collection.count() == 0


def _iter_docx_blocks(doc: DocxDocument) -> Iterator[Paragraph | Table]:
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _normalize_whitespace(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _table_to_text(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = [_normalize_whitespace(cell.text) for cell in row.cells if _normalize_whitespace(cell.text)]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _flush_section(
    documents: list[Document],
    source_path: str,
    heading_stack: list[str],
    body_parts: list[str],
    section_index: int,
) -> int:
    section_text = _normalize_whitespace("\n\n".join(body_parts))
    if not section_text:
        return section_index

    heading_path = " > ".join(heading_stack) if heading_stack else os.path.basename(source_path)
    section_title = heading_stack[-1] if heading_stack else os.path.basename(source_path)

    documents.append(
        Document(
            page_content=section_text,
            metadata={
                "source": source_path,
                "file_type": "docx",
                "section_index": section_index,
                "section_title": section_title,
                "heading_path": heading_path,
            },
        )
    )
    return section_index + 1


def _load_docx_documents(path: str) -> list[Document]:
    doc = DocxDocumentFile(path)
    source_path = _normalize_source_path(path)
    documents: list[Document] = []
    heading_stack: list[str] = []
    body_parts: list[str] = []
    section_index = 0

    for block in _iter_docx_blocks(doc):
        if isinstance(block, Paragraph):
            text = _normalize_whitespace(block.text)
            if not text:
                continue

            style_name = getattr(block.style, "name", "")
            if style_name.startswith("Heading"):
                section_index = _flush_section(documents, source_path, heading_stack, body_parts, section_index)
                body_parts = []

                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1

                heading_stack = heading_stack[: level - 1]
                heading_stack.append(text)
                continue

            body_parts.append(text)
            continue

        table_text = _table_to_text(block)
        if table_text:
            body_parts.append(table_text)

    _flush_section(documents, source_path, heading_stack, body_parts, section_index)
    return documents


def _load_json_documents(path: str) -> list[Document]:
    """Load a JSON file and convert entries into LangChain Documents."""
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    documents = []

    if isinstance(data, list):
        for i, entry in enumerate(data):
            if isinstance(entry, dict):
                content_key = next(
                    (k for k in ("content", "text", "description", "body") if k in entry),
                    None,
                )
                if content_key:
                    page_content = entry[content_key]
                    metadata = {
                        k: v
                        for k, v in entry.items()
                        if k != content_key and isinstance(v, (str, int, float, bool))
                    }
                else:
                    page_content = json.dumps(entry, ensure_ascii=False)
                    metadata = {}
                metadata["source_index"] = i
            else:
                page_content = str(entry)
                metadata = {"source_index": i}

            documents.append(Document(page_content=page_content, metadata=metadata))
    else:
        documents.append(
            Document(
                page_content=json.dumps(data, ensure_ascii=False),
                metadata={"source": _normalize_source_path(path), "file_type": "json"},
            )
        )

    return documents


def _load_source_documents(path: str) -> list[Document]:
    extension = os.path.splitext(path)[1].lower()

    if extension == ".docx":
        return _load_docx_documents(path)
    if extension == ".json":
        return _load_json_documents(path)

    raise ValueError(f"Unsupported data file type: {extension}. Supported types are .docx and .json.")


def _build_splitter() -> RecursiveCharacterTextSplitter:
    settings = get_settings()
    return RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", "؟ ", "! ", ". ", "، ", " "],
        keep_separator="end",
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )


def _build_chunk_ids(source_hash: str, chunk_count: int) -> list[str]:
    return [f"{INDEX_SCHEMA_VERSION}:{source_hash}:{index}" for index in range(chunk_count)]


def _merge_small_chunks(chunks: list[Document]) -> list[Document]:
    if not chunks:
        return chunks

    merged: list[Document] = []

    for chunk in chunks:
        if (
            merged
            and len(chunk.page_content) < MIN_CHUNK_CHARS
            and merged[-1].metadata.get("heading_path") == chunk.metadata.get("heading_path")
        ):
            merged[-1].page_content = f"{merged[-1].page_content}\n{chunk.page_content}".strip()
            continue

        merged.append(chunk)

    return merged


def _annotate_chunks(chunks: list[Document], source_path: str, source_hash: str) -> list[Document]:
    for index, chunk in enumerate(chunks):
        chunk.metadata = {
            **chunk.metadata,
            "source_path": source_path,
            "source_hash": source_hash,
            "chunk_index": index,
            "index_schema_version": INDEX_SCHEMA_VERSION,
        }
    return chunks


def _get_existing_source_chunks(vectorstore: Chroma, source_path: str) -> dict:
    return vectorstore.get(where={"source_path": source_path})


def ingest_document() -> None:
    """Load the configured data file, split it, and store embeddings in ChromaDB.

    Replaces existing chunks for the same source file and skips exact re-ingestion.
    """
    doc_path = _get_doc_path()

    if not os.path.exists(doc_path):
        logger.warning("Data file not found at %s - skipping ingestion.", doc_path)
        return

    source_path = _normalize_source_path(doc_path)
    source_hash = _compute_source_hash(source_path)
    vectorstore = get_vectorstore()

    logger.info("Loading data file: %s", source_path)
    documents = _load_source_documents(source_path)

    if not documents:
        logger.warning("No documents extracted from %s - skipping.", source_path)
        return

    splitter = _build_splitter()
    chunks = splitter.split_documents(documents)
    chunks = _merge_small_chunks(chunks)
    chunks = _annotate_chunks(chunks, source_path=source_path, source_hash=source_hash)
    chunk_ids = _build_chunk_ids(source_hash=source_hash, chunk_count=len(chunks))

    existing = _get_existing_source_chunks(vectorstore, source_path=source_path)
    existing_ids = existing.get("ids", [])
    existing_pairs = {
        (
            metadata.get("source_hash"),
            metadata.get("index_schema_version"),
        )
        for metadata in existing.get("metadatas", [])
        if metadata
    }

    if set(existing_ids) == set(chunk_ids) and existing_pairs == {(source_hash, INDEX_SCHEMA_VERSION)}:
        logger.info("Source file unchanged - skipping ingestion for %s.", source_path)
        return

    if existing_ids:
        logger.info("Replacing %d existing chunks for %s.", len(existing_ids), source_path)
        vectorstore.delete(ids=existing_ids)

    vectorstore.add_documents(chunks, ids=chunk_ids)
    logger.info("Ingested %d chunks into ChromaDB from %s.", len(chunks), source_path)
